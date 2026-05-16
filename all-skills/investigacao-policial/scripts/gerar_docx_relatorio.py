"""Generate a .docx report from a Markdown source using the PCRN/DRCC template.

Features:
  - Official state header with police badge and state coat of arms
  - SEGREDO DE JUSTICA box (optional)
  - Identification box with IP, investigados, tipificacao (optional)
  - Section headings (##) rendered as gray bordered boxes
  - GFM tables rendered as proper Word tables with header shading
  - Mermaid code blocks rendered as PNG images via mmdc (if available)
  - Times New Roman 12pt, ABNT margins, 1.5 line spacing
  - Footer with procedure reference and page numbers
"""

from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

# ── Constants ──────────────────────────────────────────────────────────────

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
FONT_SIZE_TABLE = Pt(10)
FONT_SIZE_HALF_POINTS = "24"  # 12pt in half-points
LINE_SPACING = 1.5

# ── Helpers ────────────────────────────────────────────────────────────────


def _set_cell_borders(cell, val="single", sz="6", color="000000"):
    """Set all four borders on a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        existing = tcBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(el)


def _set_cell_shading(cell, fill_color: str):
    """Set background shading on a table cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def _set_table_no_borders(table):
    """Remove all table-level borders (borders set at cell level instead)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "000000")
        el.set(qn("w:space"), "0")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_table_width(table, width_cm: float):
    """Set preferred table width in cm."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    twips_val = int(width_cm * 567)
    tblW.set(qn("w:w"), str(twips_val))
    tblW.set(qn("w:type"), "dxa")
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)


def _make_run(text: str, bold=False, font_size=None, font_name=None):
    """Create a run element with specified formatting."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        b.set(qn("w:val"), "1")
        rPr.append(b)
    if font_name:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rPr.append(rFonts)
    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), font_size)
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), font_size)
        rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


# ── Document configuration ────────────────────────────────────────────────


def configure_styles(document: Document):
    """Set up the Normal style for the document."""
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    rPr = normal._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        normal._element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = FONT_SIZE
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def configure_footer(section, footer_text: str):
    """Add footer with procedure reference and page number."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_text:
        run = paragraph.add_run(f"{footer_text}  |  P\u00e1gina ")
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.append(rFonts)
    r.append(rPr)
    fld.append(r)
    paragraph._p.append(fld)


# ── Preamble elements ─────────────────────────────────────────────────────


def add_segredo_box(document: Document):
    """Add the 'SEGREDO DE JUSTICA' box: gray background, thick border."""
    table = document.add_table(rows=1, cols=1)
    _set_table_no_borders(table)
    cell = table.cell(0, 0)
    _set_cell_borders(cell, val="single", sz="12", color="000000")
    _set_cell_shading(cell, "d9d9d9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SEGREDO DE JUSTI\u00c7A")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    document.add_paragraph()  # spacing


def add_identification_box(
    document: Document,
    ip_numero: str,
    investigados: str | None = None,
    tipificacao: list[str] | None = None,
):
    """Add the identification box with IP, investigados, tipificacao."""
    table = document.add_table(rows=1, cols=1)
    _set_table_no_borders(table)
    cell = table.cell(0, 0)
    _set_cell_borders(cell, val="single", sz="6", color="000000")

    # IP line
    p_ip = cell.paragraphs[0]
    p_ip.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p_ip.add_run("Inqu\u00e9rito Policial: ")
    run_label.bold = True
    run_label.font.name = FONT_NAME
    run_label.font.size = FONT_SIZE
    run_val = p_ip.add_run(ip_numero)
    run_val.font.name = FONT_NAME
    run_val.font.size = FONT_SIZE

    # Investigados line
    if investigados:
        p_inv = cell.add_paragraph()
        p_inv.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_label = p_inv.add_run("Investigados: ")
        run_label.bold = True
        run_label.font.name = FONT_NAME
        run_label.font.size = FONT_SIZE
        run_val = p_inv.add_run(investigados)
        run_val.font.name = FONT_NAME
        run_val.font.size = FONT_SIZE

    # Tipificacao lines
    if tipificacao:
        p_tip = cell.add_paragraph()
        p_tip.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_label = p_tip.add_run("Tipifica\u00e7\u00e3o Provis\u00f3ria:")
        run_label.bold = True
        run_label.font.name = FONT_NAME
        run_label.font.size = FONT_SIZE
        for tipo in tipificacao:
            p_t = cell.add_paragraph()
            run_t = p_t.add_run(tipo)
            run_t.font.name = FONT_NAME
            run_t.font.size = FONT_SIZE

    document.add_paragraph()  # spacing


def add_section_header_box(document: Document, title: str):
    """Add a section header in a gray bordered box."""
    table = document.add_table(rows=1, cols=1)
    _set_table_no_borders(table)
    cell = table.cell(0, 0)
    _set_cell_borders(cell, val="single", sz="12", color="000000")
    _set_cell_shading(cell, "d9d9d9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(title)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE


# ── GFM Table handling ────────────────────────────────────────────────────


def _is_table_separator(line: str) -> bool:
    """Check if a line is a GFM table separator row (|---|---|)."""
    stripped = line.strip()
    return bool(re.match(r"^\|?\s*[-:]+[-| :]*$", stripped))


def _parse_table_row(line: str) -> list[str]:
    """Parse a GFM table row into a list of cell strings."""
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [c.strip() for c in stripped.split("|")]


def add_md_table(document: Document, table_lines: list[str]):
    """Render a GFM markdown table as a properly formatted Word table."""
    if len(table_lines) < 2:
        return

    header = _parse_table_row(table_lines[0])
    # table_lines[1] is the separator row — skip
    data_rows = [_parse_table_row(ln) for ln in table_lines[2:]]

    n_cols = len(header)
    if n_cols == 0:
        return

    # Normalize all rows to same column count
    for row in data_rows:
        while len(row) < n_cols:
            row.append("")

    total_rows = 1 + len(data_rows)
    table = document.add_table(rows=total_rows, cols=n_cols)
    _set_table_no_borders(table)
    _set_table_width(table, 16.0)

    # Distribute columns evenly across 16cm
    col_width = Cm(16.0 / n_cols)
    for col in table.columns:
        col.width = col_width

    # Header row — gray background, bold, centered
    for j, cell_text in enumerate(header[:n_cols]):
        cell = table.cell(0, j)
        _set_cell_borders(cell, val="single", sz="6", color="000000")
        _set_cell_shading(cell, "d9d9d9")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run(cell_text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TABLE

    # Data rows
    for i, row in enumerate(data_rows):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.cell(i + 1, j)
            _set_cell_borders(cell, val="single", sz="6", color="000000")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = LINE_SPACING
            add_inline_formatting(p, cell_text, font_size=FONT_SIZE_TABLE)

    document.add_paragraph()  # spacing after table


# ── Mermaid rendering ─────────────────────────────────────────────────────


def _find_mmdc() -> str | None:
    """Locate the mmdc executable (supports Windows .cmd wrapper)."""
    import shutil

    for name in ("mmdc", "mmdc.cmd"):
        found = shutil.which(name)
        if found:
            return found

    # Common Windows npm global bin locations
    for candidate in [
        os.path.expanduser("~\\AppData\\Roaming\\npm\\mmdc.cmd"),
        os.path.expanduser("~\\AppData\\Roaming\\npm\\mmdc"),
    ]:
        if os.path.isfile(candidate):
            return candidate

    return None


def render_mermaid_png(mermaid_code: str, idx: int, temp_dir: str) -> str | None:
    """Render mermaid diagram code to PNG via mmdc. Returns PNG path or None."""
    mmdc = _find_mmdc()
    if not mmdc:
        print("WARNING: mmdc not found. Install with: npm install -g @mermaid-js/mermaid-cli", file=sys.stderr)
        return None

    input_file = os.path.join(temp_dir, f"mermaid_{idx}.mmd")
    output_file = os.path.join(temp_dir, f"mermaid_{idx}.png")

    try:
        with open(input_file, "w", encoding="utf-8") as f:
            f.write(mermaid_code)

        result = subprocess.run(
            [mmdc, "-i", input_file, "-o", output_file, "-b", "white", "--width", "1400"],
            capture_output=True,
            text=True,
            timeout=90,
            shell=(os.name == "nt"),
        )
        if result.returncode == 0 and os.path.exists(output_file):
            return output_file
        print(f"mmdc error (diagram {idx}): {result.stderr[:300]}", file=sys.stderr)
        return None
    except Exception as e:
        print(f"Mermaid render failed (diagram {idx}): {e}", file=sys.stderr)
        return None


# ── Markdown parsing ──────────────────────────────────────────────────────


def add_inline_formatting(paragraph, text: str, font_size=None):
    """Parse **bold** and [A PREENCHER] markers and add formatted runs."""
    if font_size is None:
        font_size = FONT_SIZE
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part
        if is_bold:
            run = paragraph.add_run(content)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = font_size
        else:
            # Split by [A PREENCHER] to apply red color
            subparts = re.split(r"(\[A PREENCHER\])", content)
            for subpart in subparts:
                if not subpart:
                    continue
                run = paragraph.add_run(subpart)
                run.font.name = FONT_NAME
                run.font.size = font_size
                if subpart == "[A PREENCHER]":
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def parse_markdown(
    document: Document,
    text: str,
    skip_ip_line: bool = False,
    temp_dir: str | None = None,
):
    """Parse Markdown text and add formatted content to the document."""
    lines = text.splitlines()
    mermaid_counter = [0]
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        if not line.strip():
            document.add_paragraph()
            continue

        # Skip IP reference line if already rendered in identification box
        if skip_ip_line and re.match(
            r"^\s*Inqu[e\u00e9]rito Policial\s+n\.\s*\d", line, re.IGNORECASE
        ):
            continue

        # H1 title
        if line.startswith("# "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line[2:].strip().upper())
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(14)
            continue

        # H2 section -> gray bordered box
        if line.startswith("## "):
            add_section_header_box(document, line[3:].strip())
            document.add_paragraph()  # spacing after box
            continue

        # H3 subsection
        if line.startswith("### "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line):
            p = document.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = LINE_SPACING
            add_inline_formatting(p, re.sub(r"^\d+\.\s+", "", line))
            continue

        # Bullet list
        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = LINE_SPACING
            add_inline_formatting(p, line[2:].strip())
            continue

        # GFM table: line starts with | and next line is a separator row
        if line.strip().startswith("|") and i < len(lines) and _is_table_separator(lines[i]):
            table_lines = [line]
            while i < len(lines) and lines[i].rstrip().strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_md_table(document, table_lines)
            continue

        # Mermaid fenced code block
        if line.strip() == "```mermaid":
            mermaid_lines = []
            while i < len(lines) and lines[i].rstrip().strip() != "```":
                mermaid_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # consume closing ```
            mermaid_code = "\n".join(mermaid_lines)
            mermaid_counter[0] += 1

            png_path = render_mermaid_png(mermaid_code, mermaid_counter[0], temp_dir) if temp_dir else None
            if png_path:
                document.add_picture(png_path, width=Cm(15.5))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Fallback: italic note
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[Diagrama {mermaid_counter[0]} — disponível no arquivo mermaid_diagramas.md]")
                run.italic = True
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
            document.add_paragraph()
            continue

        # Any other fenced code block — skip content
        if line.strip().startswith("```"):
            while i < len(lines) and not lines[i].rstrip().strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            continue

        # Horizontal rule — skip
        if re.match(r"^-{3,}$", line.strip()):
            continue

        # Regular paragraph
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(28)
        p.paragraph_format.line_spacing = LINE_SPACING
        add_inline_formatting(p, line)


# ── Main ───────────────────────────────────────────────────────────────────


def find_template(script_dir: Path) -> Path | None:
    """Look for the template in standard locations."""
    candidates = [
        script_dir / ".." / "assets" / "template_pcrn_drcc.docx",
        script_dir / "assets" / "template_pcrn_drcc.docx",
        Path.cwd() / "assets" / "template_pcrn_drcc.docx",
    ]
    for c in candidates:
        if c.resolve().exists():
            return c.resolve()
    return None


def main():
    parser = argparse.ArgumentParser(
        description="Generate a .docx report from Markdown using the PCRN/DRCC template."
    )
    parser.add_argument("--input", required=True, help="Path to Markdown source file")
    parser.add_argument("--output", required=True, help="Path to output .docx file")
    parser.add_argument("--template", default=None, help="Path to template .docx (auto-detected if omitted)")
    parser.add_argument("--segredo-justica", action="store_true", help="Add SEGREDO DE JUSTICA box at the top")
    parser.add_argument("--ip-numero", default=None, help="IP number for identification box")
    parser.add_argument("--investigados", default=None, help="Investigados string for identification box")
    parser.add_argument("--tipificacao", action="append", default=[], help="Tipificacao line (repeatable)")
    parser.add_argument("--footer-text", default=None, help="Footer text (defaults to IP reference)")
    # Backward compatibility
    parser.add_argument("--header-line", action="append", default=[])

    args = parser.parse_args()

    source = Path(args.input)
    target = Path(args.output)
    target.parent.mkdir(parents=True, exist_ok=True)

    # Locate template
    template_path = args.template
    if template_path is None:
        template_path = find_template(Path(__file__).parent)
    if template_path and Path(template_path).exists():
        document = Document(str(template_path))
        print(f"Using template: {template_path}", file=sys.stderr)
    else:
        document = Document()
        print("WARNING: Template not found — creating blank document.", file=sys.stderr)
        section = document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        sectPr = section._sectPr
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            pgMar.set(qn("w:top"), "1701")
            pgMar.set(qn("w:bottom"), "1134")
            pgMar.set(qn("w:left"), "1701")
            pgMar.set(qn("w:right"), "1134")

    configure_styles(document)

    footer_text = args.footer_text
    if footer_text is None and args.ip_numero:
        footer_text = f"IP n. {args.ip_numero}"
    configure_footer(document.sections[0], footer_text or "")

    if args.segredo_justica:
        add_segredo_box(document)

    if args.ip_numero:
        add_identification_box(
            document,
            ip_numero=args.ip_numero,
            investigados=args.investigados,
            tipificacao=args.tipificacao if args.tipificacao else None,
        )

    # Use a temp directory for Mermaid PNG files
    with tempfile.TemporaryDirectory() as temp_dir:
        md_text = source.read_text(encoding="utf-8")
        parse_markdown(document, md_text, skip_ip_line=bool(args.ip_numero), temp_dir=temp_dir)
        document.save(str(target))

    print(target.resolve())


if __name__ == "__main__":
    main()
