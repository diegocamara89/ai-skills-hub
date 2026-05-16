from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def configure_section(section):
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(0.75)


def configure_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def configure_header(section, header_lines: list[str]):
    if not header_lines:
        return
    header = section.header
    for idx, line in enumerate(header_lines):
        paragraph = header.paragraphs[0] if idx == 0 else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)


def configure_footer(section, footer_text: str):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer_text:
        run = paragraph.add_run(f"{footer_text} | Página ")
        run.font.name = "Arial"
        run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fld.append(OxmlElement("w:r"))
    paragraph._p.append(fld)


def add_inline_formatting(paragraph, text: str):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part
        run = paragraph.add_run(content)
        run.bold = is_bold
        run.font.name = "Arial"
        run.font.size = Pt(12)


def parse_markdown(document: Document, text: str):
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph()
            continue

        if line.startswith("# "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line[2:].strip().upper())
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(16)
            continue

        if line.startswith("## "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(14)
            continue

        if line.startswith("### "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(12)
            continue

        if re.match(r"^\d+\.\s+", line):
            p = document.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            add_inline_formatting(p, re.sub(r"^\d+\.\s+", "", line))
            continue

        if line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            add_inline_formatting(p, line[2:].strip())
            continue

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(28)
        p.paragraph_format.line_spacing = 1.5
        add_inline_formatting(p, line)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--header-line", action="append", default=[])
    parser.add_argument("--footer-text", default="")
    args = parser.parse_args()

    source = Path(args.input)
    target = Path(args.output)
    target.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_section(document.sections[0])
    configure_styles(document)
    configure_header(document.sections[0], args.header_line)
    configure_footer(document.sections[0], args.footer_text)
    parse_markdown(document, source.read_text(encoding="utf-8"))
    document.save(target)
    print(target.resolve())


if __name__ == "__main__":
    main()
